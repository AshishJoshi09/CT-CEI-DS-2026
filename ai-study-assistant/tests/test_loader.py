"""
Tests for app.rag.loader - these do not require an OpenAI API key since
loading/extraction has no LLM dependency.
"""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from app.rag.loader import DocumentLoadError, infer_file_type, load_document


def test_infer_file_type_pdf():
    assert infer_file_type("notes.pdf") == "pdf"


def test_infer_file_type_docx():
    assert infer_file_type("notes.docx") == "docx"


def test_infer_file_type_txt():
    assert infer_file_type("notes.txt") == "txt"


def test_infer_file_type_unsupported():
    with pytest.raises(DocumentLoadError):
        infer_file_type("notes.xyz")


def test_load_txt(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Supervised learning uses labeled data.")

    docs = load_document(file_path, "txt")

    assert len(docs) == 1
    assert "Supervised learning" in docs[0].page_content
    assert docs[0].metadata["page"] is None


def test_load_txt_empty_raises(tmp_path: Path):
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   ")

    with pytest.raises(DocumentLoadError):
        load_document(file_path, "txt")


def test_load_docx(tmp_path: Path):
    file_path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("Deadlock occurs when processes wait indefinitely.")
    document.save(str(file_path))

    docs = load_document(file_path, "docx")

    assert len(docs) == 1
    assert "Deadlock" in docs[0].page_content


def test_load_docx_empty_raises(tmp_path: Path):
    file_path = tmp_path / "empty.docx"
    document = docx.Document()
    document.save(str(file_path))

    with pytest.raises(DocumentLoadError):
        load_document(file_path, "docx")


def test_load_pdf_nonexistent_raises(tmp_path: Path):
    file_path = tmp_path / "does_not_exist.pdf"
    with pytest.raises(DocumentLoadError):
        load_document(file_path, "pdf")
