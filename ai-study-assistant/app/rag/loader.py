"""
Document loading: extracts raw text (with page numbers where applicable)
from PDF, DOCX, and TXT files and returns a list of LangChain Document
objects with consistent metadata.
"""
from __future__ import annotations

from pathlib import Path

import docx  # python-docx
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document


class DocumentLoadError(Exception):
    """Raised when a document cannot be parsed/extracted."""


def _load_pdf(file_path: Path) -> list[Document]:
    """Load a PDF, returning one Document per page with page metadata."""
    try:
        loader = PyPDFLoader(str(file_path))
        pages = loader.load()
    except Exception as exc:  # PyPDFLoader raises various low-level errors
        raise DocumentLoadError(
            f"Could not read PDF '{file_path.name}'. It may be corrupted or password-protected."
        ) from exc

    if not pages:
        raise DocumentLoadError(f"PDF '{file_path.name}' contains no extractable text.")

    docs: list[Document] = []
    for page in pages:
        text = page.page_content.strip()
        if not text:
            continue
        page_num = page.metadata.get("page", 0) + 1  # PyPDFLoader is 0-indexed
        docs.append(Document(page_content=text, metadata={"page": page_num}))

    if not docs:
        raise DocumentLoadError(
            f"PDF '{file_path.name}' has no readable text (it may be a scanned image)."
        )
    return docs


def _load_docx(file_path: Path) -> list[Document]:
    """Load a DOCX file. DOCX has no native 'page' concept, so page=None."""
    try:
        document = docx.Document(str(file_path))
    except Exception as exc:
        raise DocumentLoadError(
            f"Could not read DOCX '{file_path.name}'. The file may be corrupted."
        ) from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    if not text.strip():
        raise DocumentLoadError(f"DOCX '{file_path.name}' contains no readable text.")

    return [Document(page_content=text, metadata={"page": None})]


def _load_txt(file_path: Path) -> list[Document]:
    """Load a plain text file, trying a couple of common encodings."""
    for encoding in ("utf-8", "latin-1"):
        try:
            text = file_path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentLoadError(f"Could not decode text file '{file_path.name}'.")

    if not text.strip():
        raise DocumentLoadError(f"TXT file '{file_path.name}' is empty.")

    return [Document(page_content=text, metadata={"page": None})]


def load_document(file_path: Path, file_type: str) -> list[Document]:
    """
    Dispatch to the correct loader based on file_type ('pdf' | 'docx' | 'txt').
    Returns a list of Document objects (one per page for PDFs, one for others).
    """
    loaders = {
        "pdf": _load_pdf,
        "docx": _load_docx,
        "txt": _load_txt,
    }
    loader_fn = loaders.get(file_type)
    if loader_fn is None:
        raise DocumentLoadError(f"Unsupported file type: {file_type}")
    return loader_fn(file_path)


def infer_file_type(filename: str) -> str:
    """Map a filename's extension to an internal file_type key."""
    ext = Path(filename).suffix.lower()
    mapping = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}
    if ext not in mapping:
        raise DocumentLoadError(f"Unsupported file extension: {ext}")
    return mapping[ext]
